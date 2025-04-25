import argparse
import os
import shutil
import zipfile
import re
import difflib
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor

"""
    A class for merging visually identical spans in ODT documents.

    Motivation: DOCX documents may contain runs, which do not affect the
    appearance of the document. For instance, in some documents, a single 
    sentence may consist of multiple runs, where the only difference is that
    some runs do not have a background color, and others have a background
    color of white (which is visually identical to having no background color,
    when the page color is white, which is almost always).
    This poses problems for localization. In localization,
    visually significant spans (e.g. colors, bold, italic, footnote etc.)
    need to be transferred into the translation in correct places, but the
    presence of other, visually irrelevant spans complicates the process,
    making correct placement of spans in the translation almost impossible.
    This is also a problem in computer-assisted translation, as human
    translators find it impossible to place the spans correctly (tag soup).

    Description of functionality: This cleaner class works by processing 
    the runs in the document and merging spans that are visually identical, 
    which is defined as not differing in visually significant run properties.
    The main challenge in implementing this is that a run can get visual
    properties from multiple sources:

        1. the run properties of the run itself
        2. the style associated with the run
        3. the run properties of the paragraph
        4. the style associated with the paragraph
        5. the default style.
    
    This means that the properties need to fetched across the inheritance
    tree in order to check whether runs are visually identical.

    NOTE: When the cleaner joins visually identical runs, it will include
    all the visually identical properties as run properties. This will mean
    that specific styles associated with the run are lost (but their properties will
    remain). 

"""

#TODO: test footnotes, super/subscripts etc.

class DocxCleaner():

    def __init__(self):
        pass

    def get_style_chain(self, style):
        """Yield styles in inheritance order (closest to furthest)."""
        while style is not None:
            yield style.font
            style = style.base_style

    def get_effective_run_format(self, run, paragraph):
        """Returns a dictionary of all effective formatting properties."""
        def resolve(prop):
            # Check in priority order: run -> run style -> paragraph -> style chain
            # If run has no style, .style returns default paragraph style. This is not desirable, as
            # paragraph styles override the default style, so using the default style may override
            # formatting. Because of that, check the run._r.style (can be None) instead of .style.
            # For the paragraphs, if there is no style, .style also returns default paragraph style.
            # That is not a problem, though, since there is no style that would override it.

            # TODO: handling of rPr elements in rPr elements, those do not seem to be available
            # through the python-docx API. Are those properties actually used?
            sources = [
                run.font,
                run.style.font if run._r.style else None,
                *(self.get_style_chain(paragraph.style) if paragraph.style else [])
            ]
            for source in filter(None, sources):
                val = getattr(source, prop, None)
                if val is not None:
                    return val
            return None

        highlight = resolve('highlight_color')
        color = resolve('color')
        
        return {
            'bold': bool(resolve('bold')),
            'italic': bool(resolve('italic')),
            'underline': bool(resolve('underline')),
            'strike': bool(resolve('strike')),
            'double_strike': bool(resolve('double_strike')),
            'all_caps': bool(resolve('all_caps')),
            'small_caps': bool(resolve('small_caps')),
            'shadow': bool(resolve('shadow')),
            'outline': bool(resolve('outline')),
            'emboss': bool(resolve('emboss')),
            'imprint': bool(resolve('imprint')),
            'size': resolve('size'),
            'highlight_color': None if highlight == WD_COLOR_INDEX.WHITE else highlight,
            'color_rgb': getattr(color, 'rgb', None),
            'name': resolve('name'),
            'subscript': bool(resolve('subscript')),
            'superscript': bool(resolve('superscript'))
        }

    def merge_runs(self, paragraph):
        """Merges adjacent runs with identical formatting."""
        if len(paragraph.runs) < 2:
            return

        # Initialize with first run
        current_fmt = self.get_effective_run_format(paragraph.runs[0], paragraph)
        current_text = paragraph.runs[0].text
        merged_runs = []

        # Process subsequent runs
        for run in paragraph.runs[1:]:
            run_fmt = self.get_effective_run_format(run, paragraph)
            # Join runs that have identical runs OR runs that contain only space
            if run_fmt == current_fmt or run.text.isspace():
                current_text += run.text
            else:
                merged_runs.append((current_text, current_fmt))
                current_text = run.text
                current_fmt = run_fmt
        merged_runs.append((current_text, current_fmt))

        # Replace original runs
        p_element = paragraph._element
        for r in paragraph.runs:
            p_element.remove(r._element)

        # Add merged runs with formatting
        for text, fmt in merged_runs:
            new_run = paragraph.add_run(text)
            font = new_run.font
            
            # Set all compatible properties
            for prop, value in fmt.items():
                if value is None:
                    continue
                    
                if prop == 'color_rgb':
                    if hasattr(font, 'color'):
                        font.color.rgb = value
                elif prop == 'highlight_color':
                    font.highlight_color = value
                elif hasattr(font, prop):
                    setattr(font, prop, value)


    def iter_paragraphs_in_element(self, element):
        """Yields all paragraphs in an element (including nested tables)."""
        for child in element._element.iter():
            if isinstance(child, CT_P):
                yield Paragraph(child, element)
            elif isinstance(child, CT_Tbl):
                table = Table(child, element)
                for row in table.rows:
                    for cell in row.cells:
                        yield from self.iter_paragraphs_in_element(cell)

    def clean_docx(self, input_path, output_path):
        """Processes all paragraphs in a document."""
        doc = Document(input_path)

        # Process main document body
        for para in self.iter_paragraphs_in_element(doc):
            self.merge_runs(para)

        # Process headers and footers
        for section in doc.sections:
            for header_footer in (section.header, section.footer):
                for para in self.iter_paragraphs_in_element(header_footer):
                    self.merge_runs(para)

        doc.save(output_path)

    # This is used to validate that the conversion did not add or delete text
    # The final test is whether the texts are identical, when normalized by removing all
    # whitespace, other tests are there for debugging
    def compare_odt_files(self, file1, file2):
        """Compares the text content of two ODT files."""
        def extract_text(path):
            """Extracts plain text from an ODT file."""
            doc = Document(path)
            text = ""
            for para in self.iter_paragraphs_in_element(doc):
                text += para.text
            # Process headers and footers
            for section in doc.sections:
                for header_footer in (section.header, section.footer):
                    for para in self.iter_paragraphs_in_element(header_footer):
                        text += para.text
            return text

        text1 = extract_text(file1)
        text2 = extract_text(file2)
        
        ws_normalized_text1 = re.sub("\s+","",text1)
        ws_normalized_text2 = re.sub("\s+","",text2)

        
        if ws_normalized_text1 == ws_normalized_text2:
            print("After tag cleaning, the docs have same content without whitespaces")
        else:
            print("After tag cleaning, docs are different in addition to whitespace differences")
            diff = difflib.unified_diff(text1.split("\n"), text2.split("\n"), fromfile=file1, tofile=file2, lineterm='')
            diff_string = "\n".join(diff)
            print(diff_string)

        return ws_normalized_text1 == ws_normalized_text2

def main():
    parser = argparse.ArgumentParser(
        description="Merge adjacent runs with identical formatting in DOCX files."
    )
    parser.add_argument("input", help="Input DOCX file path")
    parser.add_argument("output", help="Output DOCX file path")
    parser.add_argument(
        "--unzip", 
        action="store_true",
        help="Unzip the output for inspection"
    )
    args = parser.parse_args()

    docx_cleaner = DocxCleaner()
    docx_cleaner.clean_docx(args.input, args.output)
    print(f"Cleaned document saved to: {args.output}")

    if args.unzip:
        unzip_dir = os.path.splitext(args.output)[0]
        if os.path.exists(unzip_dir):
            shutil.rmtree(unzip_dir)
        os.makedirs(unzip_dir)
        with zipfile.ZipFile(args.output, 'r') as zip_ref:
            zip_ref.extractall(unzip_dir)
        print(f"Unzipped output to: {unzip_dir}")

    docx_cleaner.compare_odt_files(args.input, args.output)

if __name__ == "__main__":
    main()