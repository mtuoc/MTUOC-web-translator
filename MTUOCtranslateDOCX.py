from docx import Document #pip install python-docx
from bs4 import BeautifulSoup
import sys
import requests
import random
import sys
from lxml import etree
from docx.oxml.text.paragraph import CT_P
from docx.oxml.text.run import CT_R
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def get_footnotes_part(document):
    """Get the footnotes part of the document."""
    for rel in document.part.rels.values():
        if "footnotes" in rel.target_ref:
            return rel.target_part
    return None

def get_footnotes(document):
    """Get all footnotes from the document."""
    footnotes_part = get_footnotes_part(document)
    if not footnotes_part:
        return []
    # Parse the XML content of the footnotes part
    footnotes_xml = etree.fromstring(footnotes_part.blob)
    return footnotes_xml.findall(qn('w:footnote'))

def modify_footnote_text(footnote, new_text):
    """Modify the text of a footnote."""
    for paragraph in footnote.findall(qn('w:p')):
        for run in paragraph.findall(qn('w:r')):
            text_element = run.find(qn('w:t'))
            if text_element is not None:
                text_element.text = new_text

def update_footnotes_part(footnotes_part, footnotes):
    """Update the footnotes part with modified footnotes."""
    footnotes_xml = etree.Element(qn('w:footnotes'))
    for footnote in footnotes:
        footnotes_xml.append(footnote)
    try:
        footnotes_part._blob = etree.tostring(footnotes_xml)
    except:
        pass

def get_footnote_references(paragraph):
    """Get all footnote references in a paragraph."""
    references = []
    for run in paragraph.runs:
        for element in run._element:
            if element.tag == qn('w:footnoteReference'):
                references.append((run, element.get(qn('w:id'))))
    return references

def runs_are_equivalent(run1, run2):
    """
    Check if two runs have the same formatting.
    """
    if run1.bold != run2.bold:
        return False
    if run1.italic != run2.italic:
        return False
    if run1.underline != run2.underline:
        return False
    if run1.font.color.rgb != run2.font.color.rgb:
        return False
    if run1.font.highlight_color != run2.font.highlight_color:
        return False
    if run1.font.size != run2.font.size:
        return False
    if run1.font.name != run2.font.name:
        return False
    if run1.font.all_caps != run2.font.all_caps:
        return False
    if run1.font.small_caps != run2.font.small_caps:
        return False
    if run1.font.strike != run2.font.strike:
        return False
    if run1.font.subscript != run2.font.subscript:
        return False
    if run1.font.superscript != run2.font.superscript:
        return False
    return True

def merge_equivalent_runs(paragraph):
    """
    Merge consecutive runs with the same formatting in the given paragraph.
    """
    if not paragraph.runs:
        return
    
    new_runs = []
    current_run = paragraph.runs[0]
    current_text = current_run.text

    for run in paragraph.runs[1:]:
        if runs_are_equivalent(current_run, run):
            current_text += run.text
        else:
            new_runs.append((current_run, current_text))
            current_run = run
            current_text = run.text
    
    new_runs.append((current_run, current_text))

    # Clear the paragraph's runs and add merged runs
    paragraph.clear()
    for run, text in new_runs:
        new_run = paragraph.add_run(text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.color.rgb = run.font.color.rgb
        new_run.font.highlight_color = run.font.highlight_color
        new_run.font.size = run.font.size
        new_run.font.name = run.font.name
        new_run.font.all_caps = run.font.all_caps
        new_run.font.small_caps = run.font.small_caps
        new_run.font.strike = run.font.strike
        new_run.font.subscript = run.font.subscript
        new_run.font.superscript = run.font.superscript 

def modify_run_text(run, new_text):
    """Modify the text of a run."""
    run.text = new_text

    
def extract_footnote_text(footnote):
    """Extract text from a footnote."""
    texts = []
    for paragraph in footnote.findall(qn('w:p')):
        for run in paragraph.findall(qn('w:r')):
            text_element = run.find(qn('w:t'))
            if text_element is not None:
                texts.append(text_element.text)
    return ''.join(texts)

def translate_segment_MTUOC(segment,id=101,srcLang="en-US",tgtLang="es-ES",):
    global urlMTUOC
    translation=""
    #segment=segment.strip()
    try:
        headers = {'content-type': 'application/json'}
        #params = [{ "id" : id},{ "src" : segment},{ "srcLang" : srcLang},{"tgtLang" : tgtLang}]
        params={}
        params["id"]=random.randint(0, 10000)
        params["src"]=segment
        params["srcLang"]=srcLang
        params["tgtLang"]=tgtLang
        response = requests.post(urlMTUOC, json=params, headers=headers)
        
        target = response.json()
        translation=target["tgt"]
    except:
        errormessage="Error retrieving translation from MTUOC: \n"+ str(sys.exc_info()[1])
        print("Error", errormessage, " in segment ", segment)
    return(translation)


def translate(text):
    translation=translate_segment_MTUOC(text)
    return(translation)

def parse_xml_string(html_string):
    soup = BeautifulSoup(html_string, 'html.parser')
    chunks = []
    for element in soup.contents:
        if isinstance(element, str):
            chunks.append((element, None, None))
        elif element.name in ['run']:
            
            new_tag = "<run id='"+str(element.attrs['id'])+"'>"
            chunks.append((element.text, new_tag))
    return chunks

def clean_character_formatting(run):
    if run.bold and run.style.font.bold:
        run.bold = None
    if run.italic and run.style.font.italic:
        run.italic = None
    if run.underline and run.style.font.underline:
        run.underline = None
    if run.font.color and run.font.color.rgb == run.style.font.color.rgb:
        run.font.color.rgb = None
    if run.font.size and run.font.size == run.style.font.size:
        run.font.size = None
    if run.font.name and run.font.name == run.style.font.name:
        run.font.name = None
    if run.font.highlight_color and run.font.highlight_color == run.style.font.highlight_color:
        run.font.highlight_color = None
    return(run)

def clean_paragraph_formatting(paragraph):
    p_format = paragraph.paragraph_format
    style_format = paragraph.style.paragraph_format
    
    if p_format.alignment and p_format.alignment == style_format.alignment:
        p_format.alignment = None
    if p_format.left_indent and p_format.left_indent == style_format.left_indent:
        p_format.left_indent = None
    if p_format.right_indent and p_format.right_indent == style_format.right_indent:
        p_format.right_indent = None
    if p_format.space_before and p_format.space_before == style_format.space_before:
        p_format.space_before = None
    if p_format.space_after and p_format.space_after == style_format.space_after:
        p_format.space_after = None
    if p_format.line_spacing and p_format.line_spacing == style_format.line_spacing:
        p_format.line_spacing = None

def capture_run_properties(run):
    """Capture the initial properties of a run."""
    return {
        'bold': run.bold,
        'italic': run.italic,
        'underline': run.underline,
        'rgb': run.font.color.rgb,
        'fname': run.font.name,
        'size': run.font.size,
        'highlight_color': run.font.highlight_color
    }
    
def compare_run_properties(initial_props, current_props):
    """Compare the initial properties with the current properties."""
    changes = {}
    for prop, initial_value in initial_props.items():
        current_value = current_props[prop]
        if initial_value != current_value:
            changes[prop] = (initial_value, current_value)
    return changes

def get_xml_runs(element):
    contrun=0
    clean_paragraph_formatting(element)
    runformats={}
    extracted_text=""
    previous_format=None
    if len(element.runs)>0:
        initial_properties=capture_run_properties(element.runs[0])
        runformats["initial"]=initial_properties
        for run in element.runs:
            current_properties=capture_run_properties(run)
            differences=compare_run_properties(initial_properties,current_properties)
            f={}
            text=run.text
            if len(differences)>0:
                ot="<run id='"+str(contrun)+"'>"
                runformats[ot]=current_properties
                ct="</run>"
                extracted_text +=ot+text+ct
                contrun+=1 
                run.clear()
                
            else:
                extracted_text +=text
                run.clear()
            
            previous_format=f
    return(extracted_text,runformats)

def create_element(docxelement,translated_text,runs,chunks):
    for chunk in chunks:
        text=chunk[0]
        ot=chunk[1]
        if ot==None:
            run = docxelement.add_run(text)
            formatting=runs["initial"]
            run.bold=formatting["bold"]
            run.italic=formatting["italic"]
            run.underline=formatting["underline"]
            #run.font.color=formatting["color"]
            run.font.color.rgb=formatting["rgb"]
            run.size=formatting["size"]
            run.fname=formatting["fname"]
            run.highlight_color=formatting["highlight_color"]
        else:
            run = docxelement.add_run(text)
            formatting=runs[ot]
            run.bold=formatting["bold"]
            run.italic=formatting["italic"]
            run.underline=formatting["underline"]
            #run.font.color=formatting["color"]
            run.font.color.rgb=formatting["rgb"]
            run.size=formatting["size"]
            run.fname=formatting["fname"]
            run.highlight_color=formatting["highlight_color"]
            

def is_paragraph_containing_image(paragraph):
    for run in paragraph.runs:
        if run._element.xml.find("<pic:")>-1:
            return True
    return False

def MTUOCtranslateDOCXByTagRestoring(ip,port,input_path,output_path, translate_tables=True,translate_headers=True,translate_footers=True,translate_text_boxes=True):
    """
    Reads a Word document and extracts text along with formatting tags.

    Args:
        input_path (str): Path to the Word document.

    Returns:
        str: Text content with formatting tags.
    """
    global urlMTUOC
    urlMTUOC = "http://"+ip+":"+str(port)+"/translate"
    document = Document(input_path)
    
    # REGULAR PARAGRAPHS
    for index, paragraph in enumerate(document.paragraphs):
        if is_paragraph_containing_image(paragraph):
            pass
        else:
            (extracted_text,runs)=get_xml_runs(paragraph)
            translated_text=translate(extracted_text)
            chunks=parse_xml_string(translated_text)
            create_element(paragraph,translated_text,runs,chunks)
        
    
    #TABLES
    if translate_tables:
        for table in document.tables:
            # Iterate through rows in the table
            for row in table.rows:
                # Iterate through cells in the row
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        (extracted_text,runs)=get_xml_runs(paragraph)
                        
                        translated_text=translate(extracted_text)
                        chunks=parse_xml_string(translated_text)
                        create_element(paragraph,translated_text,runs,chunks) 
                
    #HEADERS AND FOOTERS
    if translate_headers or translate_footers:
        for section in document.sections:
            if translate_headers:
                header=section.header
                for paragraph in header.paragraphs:
                    (extracted_text,runs)=get_xml_runs(paragraph)
                    translated_text=translate(extracted_text)
                    chunks=parse_xml_string(translated_text)
                    create_element(paragraph,translated_text,runs,chunks)
            if translate_footers:
                footer=section.footer
                for paragraph in footer.paragraphs:
                    (extracted_text,runs)=get_xml_runs(paragraph)
                    translated_text=translate(extracted_text)
                    chunks=parse_xml_string(translated_text)
                    create_element(paragraph,translated_text,runs,chunks)
                    
        
            
            
        document.save(output_path)

def MTUOCtranslateDOCXByChunks(ip,port,input_path,output_path, translate_tables=True,translate_headers=True,translate_footers=True,translate_text_boxes=True):
    """
    Reads a Word document and extracts text along with formatting tags.

    Args:
        input_path (str): Path to the Word document.

    Returns:
        str: Text content with formatting tags.
    """
    global urlMTUOC
    urlMTUOC = "http://"+ip+":"+str(port)+"/translate"
    document = Document(input_path)
    
    # Get all footnotes
    footnotes = get_footnotes(document)
    footnote_dict = {footnote.get(qn('w:id')): footnote for footnote in footnotes}
    
    # REGULAR PARAGRAPHS
    
    for paragraph in document.paragraphs:
        references = get_footnote_references(paragraph)
        if is_paragraph_containing_image(paragraph):
            pass
        elif references:
            for run in paragraph.runs:
                # Modify the run text while keeping footnote references intact
                if any(element.tag == qn('w:footnoteReference') for element in run._element):
                    continue  # Skip runs with footnote references
                extracted_text=run.text
                translation=translate(extracted_text)
                modify_run_text(run, translation)

            # Modify the corresponding footnotes
            for run, footnote_id in references:
                if footnote_id in footnote_dict:
                    footnotetext=extract_footnote_text(footnote_dict[footnote_id])
                    translation=translate(footnotetext)
                    modify_footnote_text(footnote_dict[footnote_id], translation)
        else:
            merge_equivalent_runs(paragraph)
            if len(paragraph.runs)>0:
                for run in paragraph.runs:
                    extracted_text=run.text
                    translation=translate(extracted_text)
                    run.text=translation
    #TABLES
    if translate_tables:
        for table in document.tables:
            # Iterate through rows in the table
            for row in table.rows:
                # Iterate through cells in the row
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        merge_equivalent_runs(paragraph)
                        if len(paragraph.runs)>0:
                            for run in paragraph.runs:
                                extracted_text=run.text
                                translation=translate(extracted_text)
                                run.text=translation
                
    #HEADERS AND FOOTERS
    if translate_headers or translate_footers:
        for section in document.sections:
            if translate_headers:
                header=section.header
                for paragraph in header.paragraphs:
                    merge_equivalent_runs(paragraph)
                    if len(paragraph.runs)>0:
                        for run in paragraph.runs:
                            extracted_text=run.text
                            translation=translate(extracted_text)
                            run.text=translation
            if translate_footers:
                footer=section.footer
                for paragraph in footer.paragraphs:
                    merge_equivalent_runs(paragraph)
                    if len(paragraph.runs)>0:
                        for run in paragraph.runs:
                            extracted_text=run.text
                            translation=translate(extracted_text)
                            run.text=translation
                    
    #FOOTNOTES



    '''
    footnotes = get_footnotes(document)
    for i, footnote in enumerate(footnotes):
        modify_footnote_text(footnote)
    '''
    # Update the footnotes part with modified footnotes
    footnotes_part = get_footnotes_part(document)
    update_footnotes_part(footnotes_part, footnotes)
    
    document.save(output_path)

def MTUOCtranslateDOCX(IP,port,input_path,output_path,strategy="ByTagRestoring"): #ByTagRestoring, byChunks
    if strategy=="byTagRestoring":
        MTUOCtranslateDOCXByTagRestoring(IP,port,input_path,output_path, translate_tables=True,translate_headers=True,translate_footers=True,translate_text_boxes=True)
    else:
        MTUOCtranslateDOCXByChunks(IP,port,input_path,output_path, translate_tables=True,translate_headers=True,translate_footers=True,translate_text_boxes=True)
        

if __name__ == "__main__":

    # python3 MTUOCtranslateDOCX.py input.docx translated.docx
    input_docx_path = sys.argv[1]
    output_docx_path = sys.argv[2]

    IP="84.88.58.132"
    #IP="192.168.1.45"
    port=8001

    MTUOCtranslateDOCX(IP,port,input_docx_path,output_docx_path,strategy="ByTagRestoring")
